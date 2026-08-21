import asyncio
import json
from pathlib import Path

import docx
import openpyxl
import pptx

from ragpoc.config import Settings
from ragpoc.db import initialize_database
from ragpoc.embeddings import FakeEmbeddingProvider
from ragpoc.ingestion import Ingestor
from ragpoc.retrieval import Retriever


def test_ingest_docx_and_search(tmp_path: Path):
    data = tmp_path / "data"
    source = tmp_path / "manual.docx"
    doc = docx.Document()
    doc.add_heading("Manual de Instalación Kubernetes", level=1)
    doc.add_paragraph("Para desplegar el cluster use el comando helm install.")
    doc.save(source)

    settings = Settings(_env_file=None, data_dir=data, allowed_upload_dir=data / "uploads")
    connection = initialize_database(settings.database_path)
    provider = FakeEmbeddingProvider()
    
    report = asyncio.run(Ingestor(connection, settings, provider).ingest(source))
    assert report["status"] == "indexed"
    assert report["chunk_count"] >= 1

    results = asyncio.run(Retriever(connection, provider).search("helm install"))
    assert len(results) > 0
    assert results[0]["filename"] == "manual.docx"
    assert "helm install" in results[0]["text"]


def test_ingest_xlsx_and_search(tmp_path: Path):
    data = tmp_path / "data"
    source = tmp_path / "inventario.xlsx"
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Servidores"
    ws.append(["Host", "IP", "Ubicación"])
    ws.append(["srv-prod-01", "10.0.0.15", "Datacenter Bogotá"])
    wb.save(source)

    settings = Settings(_env_file=None, data_dir=data, allowed_upload_dir=data / "uploads")
    connection = initialize_database(settings.database_path)
    provider = FakeEmbeddingProvider()

    report = asyncio.run(Ingestor(connection, settings, provider).ingest(source))
    assert report["status"] == "indexed"

    results = asyncio.run(Retriever(connection, provider).search("Datacenter Bogotá"))
    assert len(results) > 0
    assert results[0]["filename"] == "inventario.xlsx"
    assert "Datacenter Bogotá" in results[0]["text"]


def test_ingest_code_and_json(tmp_path: Path):
    data = tmp_path / "data"
    py_source = tmp_path / "service.py"
    py_source.write_text("def process_payment(amount: float):\n    return f'Charged {amount}'\n", encoding="utf-8")

    json_source = tmp_path / "config.json"
    json_source.write_text(json.dumps({"database_port": 5432, "env": "production"}), encoding="utf-8")

    settings = Settings(_env_file=None, data_dir=data, allowed_upload_dir=data / "uploads")
    connection = initialize_database(settings.database_path)
    provider = FakeEmbeddingProvider()
    ingestor = Ingestor(connection, settings, provider)

    rep_py = asyncio.run(ingestor.ingest(py_source))
    assert rep_py["status"] == "indexed"

    rep_json = asyncio.run(ingestor.ingest(json_source))
    assert rep_json["status"] == "indexed"

    retriever = Retriever(connection, provider)
    res_code = asyncio.run(retriever.search("process_payment"))
    assert len(res_code) > 0
    assert any(r["filename"] == "service.py" for r in res_code)

    res_json = asyncio.run(retriever.search("database_port"))
    assert len(res_json) > 0
    assert any(r["filename"] == "config.json" for r in res_json)
